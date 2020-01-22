#! python3
# Automate Boring Stuff Practice Project - Chapter 13
# customInvitations.py - Creates custom invitations from a guest list and a
# preformatted docx document
#
# NOTE: This script requires docx module to be installed.  See:
#       https://python-docx.readthedocs.io/en/latest/user/install.html
#       for more information.
#       This script also requires a preformatted document, the user can use
#       preformDoc.docx, saved in this directory.
#       This script also requires a guest list in a txt file format.  This
#       file can be downloaded here:
#       https://nostarch.com/download/
#       Automate_the_Boring_Stuff_onlinematerials_v.2.zip

import docx

# This fuction writes invites based on a given guest list text file _guestList_
# and a preformatted word document _preformDocName_ with 3 defined styles:
# myStyle1, myStyle2, and myStyle3.  An example of this document is provided in
# this directory
def writeInvites(_guestList_, _preformDocName_):
  # create docx object with the preformatted word document
  doc = docx.Document(_preformDocName_)
  # Loop through each gest in the text file
  for guests in _guestList_:
    # Write the first paragraph using the style myStyle1
    doc.add_paragraph('It would be a pleasure to have the company of', 
                      'myStyle1')
    # Write the guest name using the style myStyle2
    doc.add_paragraph(guests, 'myStyle2')
    # Write the third paragraph using the style myStyle1
    # Insert "at" as a separate run so it can be formatted independently from
    # the rest of the paragraph
    thirdParagraph = doc.add_paragraph('at')
    thirdParagraph.style = 'myStyle1'
    # Underline the text "at"
    thirdParagraph.runs[0].underline = True
    # Add more text as a separated run so it can be formatted independently
    thirdParagraph.add_run(' 11010 Memory Lane on the Evening of')
    # Write the forth paragraph using the style myStyle3
    doc.add_paragraph('April 1st', 'myStyle3')
    # Insert "at" as a separate run so it can be formatted independently from
    # the rest of the paragraph
    fifthParagraph = doc.add_paragraph('at')
    fifthParagraph.style = 'myStyle1'
    # Underline the text "at"
    fifthParagraph.runs[0].underline = True
    fifthParagraph.add_run(' 7 o' + "'" + 'clock')
    fifthParagraph.runs[1].add_break(docx.enum.text.WD_BREAK.PAGE)
  # Save document
  doc.save('invites.docx')

# This function reads the guest list text file and adds the guest to a list
def readList(_fileName_):
  readFile = open(_fileName_)
  linesList = [line.strip('\n') for line in readFile.readlines()] #This ensures
  # the removal of the '\n' new line at the end of each line
  return linesList

# User can modyfy the values below
myFileListName = 'guests.txt'
# Function call (read guest list text file and create a guests list)
myGuestList = readList(myFileListName)
myPreformDocName = 'preformDoc.docx'
# myInviteParagraphs = ['It would be a pleasure to have the company of',
#                       '',
#                       ['at','11010 Memory Lane on the Evening of'],
#                       '',
#                       'April 1st',
#                       ['at', '7 o' + "'" + 'clock']]

# Function call (create invites)
writeInvites(myGuestList, myPreformDocName)
